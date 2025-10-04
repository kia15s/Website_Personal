#!/usr/bin/env python3 """ translate_large_document.py

Script to translate large Word (.docx) or PDF files from English to Indonesian (or other languages).

Features:

Extracts text from .docx (preserving paragraphs) or PDF (per page)

Chunks text to avoid rate limits

Supports three translation backends:

1. googletrans (free, unofficial) - may be unstable


2. Google Cloud Translate API (recommended, requires API key/project)


3. DeepL API (recommended for quality, requires API key)



Outputs a translated .docx file with paragraphs in same order


Usage examples: python translate_large_document.py input.docx output_translated.docx --provider googletrans python translate_large_document.py input.pdf output_translated.docx --provider deepl --deepl-key YOUR_KEY python translate_large_document.py input.pdf output_translated.docx --provider google_api --gcloud-key PATH/TO/key.json

Dependencies: pip install python-docx PyMuPDF googletrans==4.0.0-rc1 google-cloud-translate deepl tqdm

Notes:

For very large files, this script translates chunk-by-chunk and waits between requests should you need rate limiting.

If you use the official Google Cloud Translate or DeepL API, set the appropriate key/credentials.


"""

import sys import os import argparse import time from typing import List

Text extraction

from docx import Document import fitz  # PyMuPDF

Translation options

try: from googletrans import Translator as GoogleTransTranslator HAS_GOOGLETRANS = True except Exception: HAS_GOOGLETRANS = False

try: from google.cloud import translate_v2 as translate_v2 HAS_GOOGLE_CLOUD = True except Exception: HAS_GOOGLE_CLOUD = False

try: import deepl HAS_DEEPL = True except Exception: HAS_DEEPL = False

from tqdm import tqdm

---------------------- Utilities ----------------------

def extract_from_docx(path: str) -> List[str]: doc = Document(path) paragraphs = [p.text for p in doc.paragraphs if p.text.strip() != ""] return paragraphs

def extract_from_pdf(path: str) -> List[str]: paragraphs = [] doc = fitz.open(path) for page in doc: text = page.get_text("text") # split by double newlines as paragraph boundaries parts = [p.strip() for p in text.split("\n\n") if p.strip() != ""] if not parts: # fallback split by single lines parts = [p.strip() for p in text.split("\n") if p.strip() != ""] paragraphs.extend(parts) return paragraphs

def chunk_paragraphs(paragraphs: List[str], max_chars: int = 4000) -> List[str]: """Group paragraphs into chunks under max_chars (approx safe for many APIs).""" chunks = [] cur = [] cur_len = 0 for p in paragraphs: p_len = len(p) if p_len > max_chars: # if single paragraph is huge, split it into smaller parts by sentences parts = split_long_paragraph(p, max_chars) for part in parts: if cur_len + len(part) + 1 > max_chars: if cur: chunks.append('\n\n'.join(cur)) cur = [part] cur_len = len(part) else: cur.append(part) cur_len += len(part) + 2 else: if cur_len + p_len + 2 > max_chars: if cur: chunks.append('\n\n'.join(cur)) cur = [p] cur_len = p_len else: cur.append(p) cur_len += p_len + 2 if cur: chunks.append('\n\n'.join(cur)) return chunks

def split_long_paragraph(text: str, max_chars: int) -> List[str]: sentences = text.replace('\n', ' ').split('. ') parts = [] cur = '' for s in sentences: s_clean = s.strip() if not s_clean: continue addition = (s_clean + ('.' if not s_clean.endswith('.') else '')) if len(cur) + len(addition) + 1 <= max_chars: cur = (cur + ' ' + addition).strip() else: if cur: parts.append(cur) cur = addition if cur: parts.append(cur) return parts

---------------------- Translation backends ----------------------

def translate_with_googletrans(chunks: List[str], src: str = 'en', tgt: str = 'id') -> List[str]: if not HAS_GOOGLETRANS: raise RuntimeError('googletrans package not installed or failed to import.') tr = GoogleTransTranslator() results = [] for chunk in tqdm(chunks, desc='Translating (googletrans)'): # googletrans can accept long text but may fail occasionally translated = tr.translate(chunk, src=src, dest=tgt) results.append(translated.text) time.sleep(0.2) return results

def translate_with_google_cloud(chunks: List[str], gcloud_key: str, src: str = 'en', tgt: str = 'id') -> List[str]: if not HAS_GOOGLE_CLOUD: raise RuntimeError('google-cloud-translate package not installed or failed to import.') # If user passed a JSON key path, set env var if gcloud_key: os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = gcloud_key client = translate_v2.Client() results = [] for chunk in tqdm(chunks, desc='Translating (Google Cloud)'): resp = client.translate(chunk, source_language=src, target_language=tgt) results.append(resp['translatedText']) time.sleep(0.1) return results

def translate_with_deepl(chunks: List[str], deepl_key: str, src: str = 'EN', tgt: str = 'ID') -> List[str]: if not HAS_DEEPL: raise RuntimeError('deepl package not installed or failed to import.') translator = deepl.Translator(deepl_key) results = [] for chunk in tqdm(chunks, desc='Translating (DeepL)'): resp = translator.translate_text(chunk, source_lang=src, target_lang=tgt) results.append(resp.text) time.sleep(0.1) return results

---------------------- IO ----------------------

def save_translated_docx(paragraphs_translated: List[str], output_path: str): doc = Document() for p in paragraphs_translated: doc.add_paragraph(p) doc.save(output_path)

---------------------- Main flow ----------------------

def main(): parser = argparse.ArgumentParser(description='Translate large .docx or .pdf into another language.') parser.add_argument('input', help='Input file path (.docx or .pdf)') parser.add_argument('output', help='Output docx path') parser.add_argument('--provider', choices=['googletrans', 'google_api', 'deepl'], default='googletrans') parser.add_argument('--gcloud-key', help='Path to Google Cloud JSON credentials (for google_api provider)') parser.add_argument('--deepl-key', help='DeepL API key (for deepl provider)') parser.add_argument('--src', default='en', help='Source language code (default: en)') parser.add_argument('--tgt', default='id', help='Target language code (default: id)') parser.add_argument('--chunk', type=int, default=4000, help='Max characters per chunk (default: 4000)') parser.add_argument('--sleep', type=float, default=0.0, help='Seconds to sleep between chunk translations')

args = parser.parse_args()

input_path = args.input
output_path = args.output

if not os.path.isfile(input_path):
    print('Input file not found:', input_path)
    sys.exit(1)

ext = os.path.splitext(input_path)[1].lower()
if ext == '.docx':
    print('Extracting text from DOCX...')
    paragraphs = extract_from_docx(input_path)
elif ext == '.pdf':
    print('Extracting text from PDF...')
    paragraphs = extract_from_pdf(input_path)
else:
    print('Unsupported input type. Please use .docx or .pdf')
    sys.exit(1)

if not paragraphs:
    print('No text extracted from document.')
    sys.exit(1)

print(f'Extracted {len(paragraphs)} paragraphs/blocks. Chunking...')
chunks = chunk_paragraphs(paragraphs, max_chars=args.chunk)
print(f'Created {len(chunks)} chunks for translation.')

translated_chunks = []

if args.provider == 'googletrans':
    translated_chunks = translate_with_googletrans(chunks, src=args.src, tgt=args.tgt)
elif args.provider == 'google_api':
    if not args.gcloud_key:
        print('Google Cloud provider selected but --gcloud-key not provided.')
        sys.exit(1)
    translated_chunks = translate_with_google_cloud(chunks, gcloud_key=args.gcloud_key, src=args.src, tgt=args.tgt)
elif args.provider == 'deepl':
    if not args.deepl_key:
        print('DeepL provider selected but --deepl-key not provided.')
        sys.exit(1)
    # DeepL expects uppercase language codes
    translated_chunks = translate_with_deepl(chunks, deepl_key=args.deepl_key, src=args.src.upper(), tgt=args.tgt.upper())

# Re-split chunks back into paragraphs roughly by splitting on double newlines
restored = []
for t_chunk in translated_chunks:
    parts = [p.strip() for p in t_chunk.split('\n\n') if p.strip() != '']
    restored.extend(parts)

print(f'Restored {len(restored)} translated paragraphs — saving to {output_path} ...')
save_translated_docx(restored, output_path)
print('Done!')

if name == 'main': main()